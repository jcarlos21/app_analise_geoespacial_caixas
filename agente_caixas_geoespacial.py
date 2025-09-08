import pandas as pd
import requests
from geopy.distance import geodesic
import folium
from folium.plugins import MarkerCluster
import simplekml
import os

# === IMPORTS PARA A ESTRATÉGIA DO BUFFER GEOESPACIAL (POSTES) ===
from shapely.geometry import LineString, Point
import geopandas as gpd
from pyproj import Transformer

# === NOVOS IMPORTS (necessários para as 3 camadas) ===
import math
import numpy as np

# --- IMPORTS p/ print satelital + DOCX + ZIP ---
import matplotlib.pyplot as plt
import contextily as ctx
from matplotlib.lines import Line2D
import zipfile
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
from datetime import datetime

def _utm_epsg_from_latlon(lat, lon):
    zone = int((lon + 180) // 6) + 1
    if lat >= 0:
        return 32600 + zone
    else:
        return 32700 + zone

def _build_transformers(lat, lon):
    epsg = _utm_epsg_from_latlon(lat, lon)
    to_utm = Transformer.from_crs("EPSG:4326", f"EPSG:{epsg}", always_xy=True)
    to_wgs = Transformer.from_crs(f"EPSG:{epsg}", "EPSG:4326", always_xy=True)
    return to_utm, to_wgs

def calcular_rota_osrm(coord_origem, coord_destino):
    lat1, lon1 = coord_origem  # caixa
    lat2, lon2 = coord_destino  # ponto consultado
    url = (
        f"https://router.project-osrm.org/route/v1/foot/"
        f"{lon1},{lat1};{lon2},{lat2}"
        f"?overview=full&geometries=geojson&alternatives=true&steps=false"
    )
    try:
        resp = requests.get(url)
        data = resp.json()

        rotas = data.get('routes', [])
        if not rotas:
            raise ValueError("Nenhuma rota retornada pelo OSRM.")

        melhor = min(rotas, key=lambda r: r.get('distance', float('inf')))
        rota_coords = melhor['geometry']['coordinates']
        distancia_urbana = melhor['distance']

        ponto_inicial_osrm = rota_coords[0]
        ponto_final_osrm = rota_coords[-1]

        ponto_inicial = (ponto_inicial_osrm[1], ponto_inicial_osrm[0])
        ponto_final = (ponto_final_osrm[1], ponto_final_osrm[0])

        distancia_inicial = geodesic(coord_origem, ponto_inicial).meters
        distancia_final = geodesic(coord_destino, ponto_final).meters

        distancia_total_real = distancia_inicial + distancia_urbana + distancia_final

        # Garante pontas reais (caixa e ponto)
        rota_coords.insert(0, [lon1, lat1])  # caixa real
        rota_coords.append([lon2, lat2])     # ponto real

        return rota_coords, distancia_total_real
    except Exception as e:
        print(f"[OSRM] Erro ao calcular rota: {e}. Verifique conectividade e formato das coordenadas.")
        return [], 0

# === OSRM bidirecional (A→B e B→A; escolhe a menor) ===
def calcular_rota_osrm_bidir(coord_a, coord_b):
    rota_ab, dist_ab = calcular_rota_osrm(coord_a, coord_b)
    rota_ba, dist_ba = calcular_rota_osrm(coord_b, coord_a)

    if not rota_ab and not rota_ba:
        return [], 0

    if rota_ab and (not rota_ba or dist_ab <= dist_ba):
        return rota_ab, dist_ab

    # Melhor foi B→A: reorienta para A→B
    rota_ba_reorientada = list(reversed(rota_ba))
    return rota_ba_reorientada, dist_ba

# ---------- Funções auxiliares p/ 3 camadas ----------
def _unit(vx, vy):
    n = math.hypot(vx, vy)
    if n == 0:
        return 0.0, 0.0
    return vx / n, vy / n

def _signed_offset(line_utm: LineString, p: Point):
    s = line_utm.project(p)
    on = line_utm.interpolate(s)
    s2 = min(s + 0.5, line_utm.length)
    on2 = line_utm.interpolate(s2)
    tx, ty = _unit(on2.x - on.x, on2.y - on.y)
    nx, ny = -ty, tx
    d = p.distance(on)
    vx, vy = (p.x - on.x), (p.y - on.y)
    d_signed = vx * nx + vy * ny
    return s, d, d_signed

def _angle(p0, p1, p2):
    v1 = (p0[0]-p1[0], p0[1]-p1[1])
    v2 = (p2[0]-p1[0], p2[1]-p1[1])
    n1 = math.hypot(*v1)
    n2 = math.hypot(*v2)
    if n1 == 0 or n2 == 0:
        return 0.0
    cosang = (v1[0]*v2[0] + v1[1]*v2[1]) / (n1*n2)
    cosang = max(-1.0, min(1.0, cosang))
    ang = math.degrees(math.acos(cosang))
    return ang

def gerar_rota_por_postes(
    rota_coords_osrm,
    df_postes,
    buffer_m=20,
    coord_caixa=None,
    coord_ponto=None,
    max_lateral=15.0,
    frac_inicial=0.2,
    angulo_max=60.0,
    delta_s_min=10.0,
    salto_lateral_max=0.0,
    espacamento_min=0,
    dp_tol=0,
    gap_s=70.0,
    modo_corredor="uniao",
    buffer_reta_m=None
):
    if not rota_coords_osrm:
        return [], 0.0

    # Normaliza nomes de colunas de postes
    lower_map = {c.lower(): c for c in df_postes.columns}
    lat_col = lower_map.get('latitude') or lower_map.get('lat') or lower_map.get('y')
    lon_col = lower_map.get('longitude') or lower_map.get('lon') or lower_map.get('x')
    if lat_col is None or lon_col is None:
        print("[POSTES] Colunas de latitude/longitude não foram encontradas no arquivo de postes.")
        return [], 0.0

    # Ponto médio para escolher UTM local
    mid_idx = len(rota_coords_osrm) // 2
    ref_lon, ref_lat = rota_coords_osrm[mid_idx]
    to_utm, to_wgs = _build_transformers(ref_lat, ref_lon)
    epsg = _utm_epsg_from_latlon(ref_lat, ref_lon)

    # Projeta rota OSRM para UTM e cria buffer
    line_osrm_utm = LineString([to_utm.transform(lon, lat) for lon, lat in rota_coords_osrm])
    buf_osrm = line_osrm_utm.buffer(buffer_m)

    # Linha reta e união de buffers
    if coord_caixa is not None and coord_ponto is not None:
        pA = to_utm.transform(coord_caixa[1], coord_caixa[0])
        pB = to_utm.transform(coord_ponto[1], coord_ponto[0])
        line_reta_utm = LineString([pA, pB])
        buf_reta = line_reta_utm.buffer(buffer_reta_m or buffer_m)
    else:
        line_reta_utm = None
        buf_reta = None

    if modo_corredor == "osrm":
        corredor_utm = buf_osrm
    elif modo_corredor == "reta" and buf_reta is not None:
        corredor_utm = buf_reta
    else:  # "uniao" (padrão)
        corredor_utm = buf_osrm.union(buf_reta) if buf_reta is not None else buf_osrm

    # GDF de postes em WGS84 -> reprojeta p/ UTM
    gdf_postes = gpd.GeoDataFrame(
        df_postes.copy(),
        geometry=gpd.points_from_xy(df_postes[lon_col], df_postes[lat_col]),
        crs="EPSG:4326",
    ).to_crs(f"EPSG:{epsg}")

    # Seleciona postes dentro do corredor
    postes_no_corredor = gdf_postes[gdf_postes.geometry.within(corredor_utm)]
    if postes_no_corredor.empty:
        coords_full_utm = []
        if coord_caixa is not None:
            coords_full_utm.append(to_utm.transform(coord_caixa[1], coord_caixa[0]))
        if coord_ponto is not None:
            coords_full_utm.append(to_utm.transform(coord_ponto[1], coord_ponto[0]))
        distancia_m = 0.0
        for i in range(1, len(coords_full_utm)):
            x1, y1 = coords_full_utm[i-1]
            x2, y2 = coords_full_utm[i]
            distancia_m += math.hypot(x2-x1, y2-y1)
        coords_full_lonlat = [to_wgs.transform(x, y) for (x, y) in coords_full_utm]
        return coords_full_lonlat, distancia_m

    # ---------- Cálculos s, d, d_signed ----------
    rows = []
    for _, r in postes_no_corredor.iterrows():
        p = r.geometry
        s, d, d_signed = _signed_offset(line_osrm_utm, p)
        rows.append((s, d, d_signed, p.x, p.y))
    arr = np.array(rows)  # colunas: s, d, d_signed, x, y
    arr = arr[arr[:, 0].argsort()]

    # ---------- 1) Lado dominante ----------
    n = len(arr)
    n_ini = max(1, int(frac_inicial * n))
    d_signed_ini = arr[:n_ini, 2]
    med = np.median(d_signed_ini)
    if abs(med) < 1e-6:
        pos = np.sum(arr[:, 2] > 0)
        neg = np.sum(arr[:, 2] < 0)
        med = 1.0 if pos >= neg else -1.0
    sign_dom = 1.0 if med >= 0 else -1.0

    mask_lado = np.sign(arr[:, 2]) == sign_dom
    mask_lateral = arr[:, 1] <= max_lateral
    base_arr = arr[mask_lado & mask_lateral]

    # ---------- 3a) Tolerância a gaps ----------
    if len(base_arr) >= 2:
        s_base = base_arr[:, 0]
        gaps_idx = np.where(np.diff(s_base) > gap_s)[0]
        for gi in gaps_idx:
            s_a, s_b = s_base[gi], s_base[gi+1]
            relax = max_lateral * 1.2
            win_mask = (arr[:, 0] > s_a) & (arr[:, 0] < s_b) & (arr[:, 1] <= relax)
            candidatos = arr[win_mask]
            if len(candidatos) > 0:
                base_arr = np.vstack([base_arr, candidatos])
        base_arr = base_arr[base_arr[:, 0].argsort()]

    # ---------- 2) Regras anti-serrilha ----------
    seq = base_arr.tolist()

    def _filtra_angulos(seq_pts, ang_max=60.0, ds_min=10.0):
        if len(seq_pts) < 3:
            return seq_pts
        res = [seq_pts[0]]
        for i in range(1, len(seq_pts)-1):
            prev = res[-1]
            cur = seq_pts[i]
            nxt = seq_pts[i+1]
            p0 = (prev[3], prev[4])
            p1 = (cur[3], cur[4])
            p2 = (nxt[3], nxt[4])
            ang = _angle(p0, p1, p2)
            ds = nxt[0] - prev[0]
            if ang > ang_max and ds < ds_min:
                continue
            res.append(cur)
        res.append(seq_pts[-1])
        return res

    def _filtra_salto_lateral(seq_pts, salto_max=0.0, ds_min=10.0):
        if len(seq_pts) < 2:
            return seq_pts
        res = [seq_pts[0]]
        for i in range(1, len(seq_pts)):
            prev = res[-1]
            cur = seq_pts[i]
            salto = abs(cur[2] - prev[2])
            ds = cur[0] - prev[0]
            if salto > salto_max and ds < ds_min:
                continue
            res.append(cur)
        return res

    def _thinning(seq_pts, min_s=0):
        if len(seq_pts) < 2:
            return seq_pts
        res = [seq_pts[0]]
        last_s = seq_pts[0][0]
        for i in range(1, len(seq_pts)-1):
            if seq_pts[i][0] - last_s >= min_s:
                res.append(seq_pts[i])
                last_s = seq_pts[i][0]
        res.append(seq_pts[-1])
        return res

    seq = _filtra_angulos(seq, ang_max=angulo_max, ds_min=delta_s_min)
    seq = _filtra_salto_lateral(seq, salto_max=salto_lateral_max, ds_min=delta_s_min)
    seq = _thinning(seq, min_s=espacamento_min)

    coords_full_utm = []
    if coord_caixa is not None:
        coords_full_utm.append(to_utm.transform(coord_caixa[1], coord_caixa[0]))
    for s, d, dsig, x, y in seq:
        coords_full_utm.append((x, y))
    if coord_ponto is not None:
        coords_full_utm.append(to_utm.transform(coord_ponto[1], coord_ponto[0]))

    if len(coords_full_utm) >= 3:
        line = LineString(coords_full_utm)
        line_simpl = line.simplify(dp_tol, preserve_topology=False)
        coords_full_utm = list(line_simpl.coords)

    distancia_m = 0.0
    for i in range(1, len(coords_full_utm)):
        x1, y1 = coords_full_utm[i-1]
        x2, y2 = coords_full_utm[i]
        distancia_m += math.hypot(x2-x1, y2-y1)

    coords_full_lonlat = [to_wgs.transform(x, y) for (x, y) in coords_full_utm]
    return coords_full_lonlat, distancia_m

def gerar_kmz(nome_base, rota_coords, ponto_consultado, caixa_mais_proxima, caixa_mais_proxima_nome, nome_ponto_referencia=None):
    kml = simplekml.Kml()
    linha = kml.newlinestring(name="Rota entre ponto e caixa")
    linha.coords = rota_coords
    linha.style.linestyle.color = simplekml.Color.red
    linha.style.linestyle.width = 4

    nome_ponto_kmz = nome_ponto_referencia if nome_ponto_referencia else "Ponto de Referência"

    ponto_ref = kml.newpoint(
        name=nome_ponto_kmz,
        coords=[ponto_consultado],
        description=f"Localização consultada: {ponto_consultado}",
    )
    ponto_ref.style.iconstyle.icon.href = "http://maps.google.com/mapfiles/kml/pushpin/ltblu-pushpin.png"

    caixa_ponto = kml.newpoint(
        name=f"{caixa_mais_proxima_nome}",
        coords=[caixa_mais_proxima],
        description=f"Coordenadas: {caixa_mais_proxima}",
    )
    caixa_ponto.style.iconstyle.color = simplekml.Color.white
    caixa_ponto.style.iconstyle.scale = 1.2
    caixa_ponto.style.iconstyle.icon.href = "http://maps.google.com/mapfiles/kml/shapes/donut.png"

    os.makedirs("saida_kmz", exist_ok=True)
    kml_path = os.path.join("saida_kmz", f"{nome_base}.kmz")
    kml.savekmz(kml_path)
    return kml_path

# --------- normalização robusta do df_caixas ---------
def _normalize_caixas_df(df: pd.DataFrame) -> pd.DataFrame:
    cols = {c.lower().strip(): c for c in df.columns}
    def pick(*names):
        for n in names:
            if n in cols:
                return cols[n]
        return None
    mapping = {}
    lat = pick('latitude','lat','y')
    lon = pick('longitude','lon','x')
    cidade = pick('cidade','município','municipio')
    estado = pick('estado','uf','sigla uf')
    pasta = pick('pasta','categoria')
    sigla = pick('sigla','identificador','id','nome','caixa','id_caixa')
    if lat: mapping[lat] = 'Latitude'
    if lon: mapping[lon] = 'Longitude'
    if cidade: mapping[cidade] = 'Cidade'
    if estado: mapping[estado] = 'Estado'
    if pasta: mapping[pasta] = 'Pasta'
    if sigla: mapping[sigla] = 'Sigla'
    df2 = df.rename(columns=mapping).copy()

    # validações mínimas
    for req in ['Latitude', 'Longitude']:
        if req not in df2.columns:
            raise KeyError(f"Coluna obrigatória ausente no arquivo de caixas: '{req}'.")
    return df2

def analisar_distancia_entre_pontos(
    df_pontos, df_caixas, limite_fibra=350, df_postes=None, buffer_postes_m=None,
    usar_postes=False, k_top=3, progress_cb=None  # progress_cb opcional
):
    # normaliza df_pontos
    df_pontos.columns = df_pontos.columns.str.strip().str.upper()
    df_pontos.rename(columns={
        'NOME': 'Nome',
        'NOME MUNICÍPIO': 'Cidade',
        'SIGLA UF': 'Estado',
        'LATITUDE': 'LATITUDE',
        'LONGITUDE': 'LONGITUDE'
    }, inplace=True)

    # normaliza df_caixas
    df_caixas = _normalize_caixas_df(df_caixas)

    resultados = []

    if df_caixas.empty:
        return pd.DataFrame([])

    K_TOP = max(1, int(k_top))

    total_pontos = len(df_pontos)
    processados = 0

    for index, ponto in df_pontos.iterrows():
        coord_ponto = (ponto['LATITUDE'], ponto['LONGITUDE'])

        # Pré-seleção por geodésica: pega as K caixas mais próximas
        def _dist_geo_row(row):
            return geodesic(coord_ponto, (row['Latitude'], row['Longitude'])).meters

        df_tmp = df_caixas.copy()
        df_tmp['__dist_geo__'] = df_tmp.apply(_dist_geo_row, axis=1)
        candidatas = df_tmp.nsmallest(K_TOP, '__dist_geo__')

        # Entre as K candidatas, escolhe a menor rota real (OSRM) usando BIDIR
        melhor = None
        for _, cand in candidatas.iterrows():
            coord_caixa_cand = (cand['Latitude'], cand['Longitude'])

            rota_coords_cand, distancia_real_cand = calcular_rota_osrm_bidir(coord_caixa_cand, coord_ponto)

            if not rota_coords_cand:
                # fallback: usa geodésica dessa candidata
                distancia_real_cand = cand['__dist_geo__']

            if (melhor is None) or (distancia_real_cand < melhor['dist']):
                melhor = {
                    'caixa': cand,
                    'rota': rota_coords_cand,
                    'dist': distancia_real_cand
                }

        if melhor is None:
            processados += 1
            if progress_cb:
                progress_cb(processados, total_pontos)
            continue

        caixa_proxima = melhor['caixa']
        coord_caixa_proxima = (caixa_proxima['Latitude'], caixa_proxima['Longitude'])
        rota_coords = melhor['rota']
        distancia_real = melhor['dist']

        # Se por algum motivo a rota não veio, usa geodésica
        if not rota_coords:
            distancia_real = geodesic(coord_ponto, coord_caixa_proxima).meters

        distancia_metros = round(distancia_real, 2)
        tipo_cabo = "Drop" if distancia_metros <= 250 else "Auto Sustentado"

        ponto_coords = (ponto['LONGITUDE'], ponto['LATITUDE'])
        caixa_coords = (caixa_proxima['Longitude'], caixa_proxima['Latitude'])
        nome_ponto = ponto.get('Nome', '').strip() if isinstance(ponto.get('Nome', ''), str) else ''
        nome_caixa = caixa_proxima.get("Sigla", "")

        rota_postes_coords = []
        dist_postes_m = None
        postes_pts_coords = []
        if usar_postes and rota_coords and df_postes is not None and buffer_postes_m:
            try:
                # Corredor híbrido já aplicado dentro da função
                rota_postes_coords, dist_postes_m = gerar_rota_por_postes(
                    rota_coords, df_postes, buffer_m=float(buffer_postes_m),
                    coord_caixa=coord_caixa_proxima, coord_ponto=coord_ponto
                )
                # também guardamos os pontos de postes usados para desenhar (quando solicitado)
                if rota_postes_coords:
                    postes_pts_coords = rota_postes_coords[1:-1]  # internos
            except Exception as e:
                print(f"[POSTES] Falha ao gerar rota por postes: {e}")

        resultados.append({
            'Nome do Ponto de Referência': nome_ponto,
            'Cidade do Ponto': ponto.get('Cidade', ''),
            'Estado do Ponto': ponto.get('Estado', ''),
            'Localização do Ponto': f"{ponto['LATITUDE']}, {ponto['LONGITUDE']}",
            'Localização da Caixa': f"{caixa_proxima['Latitude']}, {caixa_proxima['Longitude']}",
            'Identificador': nome_caixa,
            'Cidade': caixa_proxima.get('Cidade', ''),
            'Estado': caixa_proxima.get('Estado', ''),
            'Categoria': caixa_proxima.get('Pasta', ''),
            'Distância da Rota (m)': distancia_metros,
            'Tipo de Cabo': tipo_cabo,
            'Viabilidade': 'Viável' if distancia_metros < limite_fibra else '',
            'Distância via Postes (m)': round(dist_postes_m, 2) if dist_postes_m else None,
            'Rota via Postes?': 'Sim' if rota_postes_coords else 'Não',
            'rota_postes_coords': rota_postes_coords,     # para print/KMZ
            'dist_postes_m': dist_postes_m,               # para print/KMZ
            'rota_osrm_coords': rota_coords or [],        # para print
            'postes_pts_coords': postes_pts_coords,       # pontos dos postes (se houver)
            'Download da Rota (KMZ)': "[Gerado em KMZ único]" if rota_coords else ""
        })

        processados += 1
        if progress_cb:
            progress_cb(processados, total_pontos)

    # --- Gera o KMZ único com as rotas e postes disponíveis
    if resultados:
        nome_arquivo_unico = "rotas_completas"
        kml = simplekml.Kml()
        for l in resultados:
            # Rota OSRM
            lat_cx, lon_cx = map(float, l['Localização da Caixa'].split(', '))
            lat_pt, lon_pt = map(float, l['Localização do Ponto'].split(', '))
            rota_osrm_coords, _ = calcular_rota_osrm_bidir((lat_cx, lon_cx), (lat_pt, lon_pt))

            pasta = kml.newfolder(name=l['Nome do Ponto de Referência'] or "Ponto")

            if l.get('Viabilidade','') != "Viável":
                cor_linha = simplekml.Color.red
            elif l.get('Tipo de Cabo','') == "Drop":
                cor_linha = simplekml.Color.rgb(0, 0, 255)
            elif l.get('Tipo de Cabo','') == "Auto Sustentado":
                cor_linha = simplekml.Color.rgb(0, 255, 0)
            else:
                cor_linha = simplekml.Color.gray

            # rota OSRM
            if rota_osrm_coords:
                linha = pasta.newlinestring(name=f"Rota OSRM - {l['Nome do Ponto de Referência'] or 'Ponto'}")
                linha.coords = rota_osrm_coords
                linha.style.linestyle.color = cor_linha
                linha.style.linestyle.width = 4

            # rota por postes
            rota_postes_coords = l.get("rota_postes_coords") or []
            if rota_postes_coords:
                lp = pasta.newlinestring(name=f"Trajeto por Postes - {l['Nome do Ponto de Referência'] or 'Ponto'}")
                lp.coords = rota_postes_coords
                lp.style.linestyle.color = simplekml.Color.changealphaint(160, simplekml.Color.cyan)
                lp.style.linestyle.width = 3
                if l.get("dist_postes_m"):
                    lp.description = f"Rota por postes (~{round(l['dist_postes_m'], 2)} m)"

                # marcar postes ao longo da rota
                for i, (lon, lat) in enumerate(rota_postes_coords[1:-1], start=1):
                    p = pasta.newpoint(name=f"Poste {i}", coords=[(lon, lat)])
                    p.style.iconstyle.icon.href = "http://maps.google.com/mapfiles/kml/paddle/P.png"

            # ponto e caixa
            ponto_ref = pasta.newpoint(
                name=l['Nome do Ponto de Referência'] or "Ponto",
                coords=[(lon_pt, lat_pt)],
                description=f"Localização consultada: {(lon_pt, lat_pt)}",
            )
            ponto_ref.style.iconstyle.icon.href = "http://maps.google.com/mapfiles/kml/pushpin/ltblu-pushpin.png"

            caixa_ponto = pasta.newpoint(
                name=l['Identificador'],
                coords=[(lon_cx, lat_cx)],
                description=f"Coordenadas: {(lon_cx, lat_cx)}",
            )
            caixa_ponto.style.iconstyle.color = simplekml.Color.white
            caixa_ponto.style.iconstyle.scale = 1.2
            caixa_ponto.style.iconstyle.icon.href = "http://maps.google.com/mapfiles/kml/shapes/donut.png"

        os.makedirs("saida_kmz", exist_ok=True)
        kml_path = os.path.join("saida_kmz", f"{nome_arquivo_unico}.kmz")
        kml.savekmz(kml_path)
        for r in resultados:
            r["Download da Rota (KMZ)"] = kml_path

    return pd.DataFrame(resultados)

def gerar_mapa_interativo(df_resultados, caminho_html):
    mapa = folium.Map(location=[-5.8, -36.6], zoom_start=8)
    marcadores = MarkerCluster().add_to(mapa)

    for _, linha in df_resultados.iterrows():
        lat_ponto, lon_ponto = map(float, linha['Localização do Ponto'].split(', '))
        lat_caixa, lon_caixa = map(float, linha['Localização da Caixa'].split(', '))

        # usa cálculo bidirecional também no mapa
        rota_coords, _ = calcular_rota_osrm_bidir((lat_caixa, lon_caixa), (lat_ponto, lon_ponto))
        if rota_coords:
            rota_convertida = [(lat, lon) for lon, lat in rota_coords]
            folium.PolyLine(
                locations=rota_convertida,
                color='red', weight=4,
                tooltip=f"Rota OSRM: {linha['Distância da Rota (m)']} metros."
            ).add_to(mapa)
        else:
            folium.PolyLine(
                locations=[[lat_ponto, lon_ponto], [lat_caixa, lon_caixa]],
                color='gray', weight=2, dash_array="5,5",
                tooltip="Rota linear (falha OSRM)"
            ).add_to(mapa)

        folium.Marker(
            location=[lat_caixa, lon_caixa],
            tooltip=f"{linha['Identificador']}",
            icon=folium.Icon(color='green', icon='hdd', prefix='fa')
        ).add_to(marcadores)

        folium.Marker(
            location=[lat_ponto, lon_ponto],
            tooltip=f"{linha['Nome do Ponto de Referência']}",
            icon=folium.Icon(color='blue', icon='info-sign')
        ).add_to(marcadores)

    mapa.save(caminho_html)
    return caminho_html

# =========================
# Classificação de postes e BOM
# =========================
def classificar_postes_por_angulo_e_regra(rota_postes_coords, ang_inf=165.0, ang_sup=195.0, max_passantes=4):
    """
    - Primeiro poste após a caixa é 'encab';
    - Se o ângulo no poste (p0,p1,p2) fugir do corredor [ang_inf, ang_sup] OU se já houve 'max_passantes'
      desde o último encabeçado, o poste vira 'encab'; senão é 'pass'.
    - Retorna também contagem de materiais (sem dielétrico).
    """
    to_3857 = Transformer.from_crs("EPSG:4326", "EPSG:3857", always_xy=True)
    coords_xy = [to_3857.transform(lon, lat) for (lon, lat) in rota_postes_coords]

    postes_info = []
    qtd_enc, qtd_pas = 0, 0
    passantes_desde_encab = 0

    for i in range(1, len(coords_xy) - 1):
        lon, lat = rota_postes_coords[i]
        if i == 1:
            tipo = "encab"
            passantes_desde_encab = 0
        else:
            p0, p1, p2 = coords_xy[i - 1], coords_xy[i], coords_xy[i + 1]
            ang = _angle(p0, p1, p2)
            if (ang < ang_inf) or (ang > ang_sup) or (passantes_desde_encab >= max_passantes):
                tipo = "encab"
                passantes_desde_encab = 0
            else:
                tipo = "pass"
                passantes_desde_encab += 1

        postes_info.append({"idx": i, "lon": lon, "lat": lat, "tipo": tipo})
        if tipo == "encab":
            qtd_enc += 1
        else:
            qtd_pas += 1

    # Materiais (sem dielétrico)
    qtd_supa = 2 * qtd_enc + 1 * qtd_pas
    qtd_bap  = 1 * qtd_enc + 1 * qtd_pas
    qtd_alca = 2 * qtd_enc

    return {
        "postes": postes_info,
        "QtdPosteEnc": qtd_enc,
        "QtdPostePas": qtd_pas,
        "QtdSUPA": qtd_supa,
        "QtdBAP": qtd_bap,
        "QtdALCA": qtd_alca
    }

# =========================
# PRINT SATELITAL + DOCX + ZIP
# =========================
def _plot_print_satelital_esri(
    ponto_latlon, caixa_latlon, rota_osrm_coords,
    rota_postes_coords=None, incluir_postes=False,
    postes_coords=None, out_png_path=None,
    padding_m=60, width_px=1800, height_px=1200, dpi=180,
    show_osrm=True, show_postes=True,
    zoom_override=None, overlay_labels=False
):
    """
    Gera um PNG satelital (Esri.WorldImagery) com:
    - marcador da CAIXA, marcador do PONTO,
    - rota OSRM (opcional via show_osrm),
    - rota por postes (opcional via show_postes),
    - (opcional) marcadores dos postes (incluir_postes + postes_coords).
    Parâmetros extras:
      - zoom_override: força nível de zoom (int) no basemap.
      - overlay_labels: adiciona camada de rótulos/transport sobre o satélite.
    """
    # fallback: se nenhuma rota for selecionada, desenha a OSRM (quando existir) ou reta
    if not show_osrm and not show_postes:
        show_osrm = True

    # Se faltar OSRM e ela for pedida, usa reta caixa↔ponto
    if show_osrm and not rota_osrm_coords:
        rota_osrm_coords = [(ponto_latlon[1], ponto_latlon[0]), (caixa_latlon[1], caixa_latlon[0])]

    # --- GeoDataFrames em WGS84
    layers = []
    if show_osrm and rota_osrm_coords:
        gdf_osrm = gpd.GeoDataFrame(geometry=[LineString([(x, y) for x, y in rota_osrm_coords])], crs="EPSG:4326")
        layers.append(("osrm", gdf_osrm))
    gdf_pt = gpd.GeoDataFrame(geometry=[Point(ponto_latlon[1], ponto_latlon[0])], crs="EPSG:4326")
    gdf_cx = gpd.GeoDataFrame(geometry=[Point(caixa_latlon[1], caixa_latlon[0])], crs="EPSG:4326")
    layers.extend([("pt", gdf_pt), ("cx", gdf_cx)])

    if show_postes and rota_postes_coords:
        gdf_postes_line = gpd.GeoDataFrame(geometry=[LineString([(x, y) for x, y in rota_postes_coords])], crs="EPSG:4326")
        layers.append(("postes_line", gdf_postes_line))

    if incluir_postes and postes_coords:
        gdf_postes_pts = gpd.GeoDataFrame(geometry=[Point(lon, lat) for (lon, lat) in postes_coords], crs="EPSG:4326")
        layers.append(("postes_pts", gdf_postes_pts))

    # --- reprojeta tudo para Web Mercator (3857)
    layers_3857 = [(name, gdf.to_crs("EPSG:3857")) for name, gdf in layers]
    dict_3857 = {name: gdf for name, gdf in layers_3857}

    # --- bbox com padding
    total = None
    for _, g in layers_3857:
        total = g if total is None else pd.concat([total, g])
    minx, miny, maxx, maxy = total.total_bounds
    pad = padding_m
    minx -= pad; miny -= pad; maxx += pad; maxy += pad

    # --- figura (SEM MARGENS)
    fig_w_in = width_px / dpi
    fig_h_in = height_px / dpi
    fig, ax = plt.subplots(1, 1, figsize=(fig_w_in, fig_h_in), dpi=dpi)
    fig.subplots_adjust(left=0, right=1, top=1, bottom=0)
    ax.set_position([0, 0, 1, 1])

    # --- basemap satelital
    ax.set_xlim(minx, maxx); ax.set_ylim(miny, maxy)
    try:
        if zoom_override is not None:
            z = int(zoom_override)
        else:
            span_m = max(maxx - minx, maxy - miny)
            if span_m <= 150:      z = 19
            elif span_m <= 300:    z = 18
            elif span_m <= 800:    z = 17
            elif span_m <= 2000:   z = 16
            else:                  z = 15

        ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, crs="EPSG:3857", zoom=z)
        if overlay_labels:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldTransportation, crs="EPSG:3857", zoom=z)
    except Exception:
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImageryClarity, crs="EPSG:3857", zoom=17)
        except Exception:
            ax.set_facecolor("white")
            ax.text(0.01, 0.01, "Satélite indisponível nesta execução", transform=ax.transAxes, fontsize=8, color="gray")

    # --- camadas (desenho)
    if dict_3857.get("osrm") is not None:
        dict_3857["osrm"].plot(ax=ax, color=None, edgecolor="tab:red", linewidth=3, zorder=5)
    if dict_3857.get("postes_line") is not None:
        dict_3857["postes_line"].plot(ax=ax, color=None, edgecolor="tab:cyan", linewidth=2.5, linestyle="--", zorder=6)

    dict_3857["cx"].plot(ax=ax, marker="s", markersize=70, edgecolor="black", facecolor="yellow", zorder=7)
    dict_3857["pt"].plot(ax=ax, marker="o", markersize=70, edgecolor="black", facecolor="lime", zorder=7)

    if dict_3857.get("postes_pts") is not None:
        dict_3857["postes_pts"].plot(ax=ax, marker=".", markersize=12, color="cyan", alpha=0.9, zorder=6)

    # --- legenda dentro do mapa
    legend_elems = []
    if dict_3857.get("osrm") is not None:
        legend_elems.append(Line2D([0],[0], color="tab:red", lw=3, label="Rota OSRM"))
    if dict_3857.get("postes_line") is not None:
        legend_elems.append(Line2D([0],[0], color="tab:cyan", lw=2.5, ls="--", label="Rota por Postes"))
    legend_elems.extend([
        Line2D([0],[0], marker="s", color="black", markerfacecolor="yellow", markersize=10, lw=0, label="Caixa"),
        Line2D([0],[0], marker="o", color="black", markerfacecolor="lime", markersize=10, lw=0, label="Ponto"),
    ])
    ax.legend(handles=legend_elems, loc="lower right", fontsize=8, framealpha=0.7,
              borderpad=0.3, handlelength=2, bbox_to_anchor=(0.98, 0.02), bbox_transform=ax.transAxes)

    ax.axis("off")
    os.makedirs(os.path.dirname(out_png_path), exist_ok=True)
    plt.savefig(out_png_path, dpi=dpi, bbox_inches=None, pad_inches=0)
    plt.close(fig)
    return out_png_path

# --- utilitário: substituição de placeholders mantendo formato do modelo
def _replace_placeholders_in_document(doc: Document, mapping: dict):
    """Substitui placeholders {{CHAVE}} em parágrafos e tabelas, preservando o formato dos runs do modelo."""
    def repl_in_runs(runs, key, val):
        for run in runs:
            if key in run.text:
                run.text = run.text.replace(key, val)

    for p in doc.paragraphs:
        for k, v in mapping.items():
            repl_in_runs(p.runs, f"{{{{{k}}}}}", v)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in mapping.items():
                        repl_in_runs(p.runs, f"{{{{{k}}}}}", v)

def _insert_map_at_placeholder(doc: Document, placeholder: str, img_path: str, width_inches=6.5):
    token = f"{{{{{placeholder}}}}}"

    for p in doc.paragraphs:
        full = "".join(run.text for run in p.runs)
        if token in full:
            for run in list(p.runs):
                run.text = ""
            run = p.add_run()
            run.add_picture(img_path, width=Inches(width_inches))
            return True

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = "".join(run.text for run in p.runs)
                    if token in full:
                        for run in list(p.runs):
                            run.text = ""
                        run = p.add_run()
                        run.add_picture(img_path, width=Inches(width_inches))
                        return True

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    return False

def _criar_docx_descritivo(simples_dict, png_path, out_docx_path, bom_dict=None, template_path="Modelo de Descritivo.docx"):
    """
    Gera o DOCX usando o template com placeholders.
    Placeholders aceitos (exemplos): {{PONTO}}, {{COORD_PONTO}}, {{CAIXA}}, {{COORD_CAIXA}},
      {{CIDADE_UF}}, {{DATA}}, {{DIST_OSRM}}, {{DIST_POSTES}}, {{TIPO_CABO}}, {{VIABILIDADE}},
      {{POSTES_ENC}}, {{POSTES_PAS}}, {{QTD_SUPA}}, {{QTD_BAP}}, {{QTD_ALCA}}, {{MAPA}}.
    """
    # prepara dados
    dist_osrm = simples_dict.get('Distância da Rota (m)')
    dist_postes = simples_dict.get('Distância via Postes (m)')
    tipo = simples_dict.get('Tipo de Cabo') or ""
    viab = simples_dict.get('Viabilidade', '') or ""
    cidadeuf = f"{simples_dict.get('Cidade','')}/{simples_dict.get('Estado','')}"
    qtd_enc = bom_dict.get("QtdPosteEnc", 0) if bom_dict else 0
    qtd_pas = bom_dict.get("QtdPostePas", 0) if bom_dict else 0
    qtd_supa = bom_dict.get("QtdSUPA", 0) if bom_dict else 0
    qtd_bap  = bom_dict.get("QtdBAP", 0) if bom_dict else 0
    qtd_alca = bom_dict.get("QtdALCA", 0) if bom_dict else 0

    mapping = {
        "PONTO": simples_dict.get('Nome do Ponto de Referência') or "—",
        "COORD_PONTO": simples_dict.get('Localização do Ponto') or "—",
        "CAIXA": simples_dict.get('Identificador') or "—",
        "COORD_CAIXA": simples_dict.get('Localização da Caixa') or "—",
        "CIDADE_UF": cidadeuf,
        "DATA": datetime.now().strftime('%d/%m/%Y %H:%M'),
        "DIST_OSRM": f"{dist_osrm} m" if dist_osrm is not None else "—",
        "DIST_POSTES": f"{dist_postes} m" if dist_postes else "—",
        "TIPO_CABO": tipo,
        "VIABILIDADE": viab,
        "POSTES_ENC": str(qtd_enc),
        "POSTES_PAS": str(qtd_pas),
        # >>> NOVO: placeholders de materiais
        "QTD_SUPA": str(qtd_supa),
        "QTD_BAP": str(qtd_bap),
        "QTD_ALCA": str(qtd_alca),
    }

    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()  # fallback

    _replace_placeholders_in_document(doc, mapping)
    _insert_map_at_placeholder(doc, "MAPA", png_path, width_inches=6.5)

    os.makedirs(os.path.dirname(out_docx_path), exist_ok=True)
    doc.save(out_docx_path)
    return out_docx_path

def gerar_descritivos_zip(df_resultados, incluir_rota_postes=True, incluir_postes_pts=False, rota_no_print="Ambas", qualidade_print="Padrão"):
    """
    Para cada linha (ponto) do df_resultados:
    - gera print satelital (PNG) com rotas conforme 'rota_no_print',
    - gera DOCX a partir do modelo (inclui contagem de postes e materiais),
    - compacta todos em um ZIP e retorna o caminho.

    qualidade_print: "Padrão" | "Alta" | "Ultra"
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_out = os.path.join("out", f"descritivos_{ts}")
    pasta_maps = os.path.join(base_out, "maps")
    pasta_docs = os.path.join(base_out, "docs")
    os.makedirs(pasta_maps, exist_ok=True)
    os.makedirs(pasta_docs, exist_ok=True)

    # Qualidade → presets
    qp = (qualidade_print or "Padrão").strip().lower()
    if qp == "alta":
        dpi = 220; width_px = 2400; height_px = 1600; zoom_override = 18; overlay_labels = True
    elif qp == "ultra":
        dpi = 300; width_px = 3000; height_px = 2000; zoom_override = 19; overlay_labels = True
    else:  # Padrão
        dpi = 180; width_px = 1800; height_px = 1200; zoom_override = None; overlay_labels = False

    # Mapeia escolha do usuário para flags de desenho
    rota_no_print_norm = (rota_no_print or "Ambas").strip().lower()
    show_osrm = rota_no_print_norm in ("osrm", "ambas")
    show_postes = rota_no_print_norm in ("postes", "ambas")

    docx_paths = []

    for _, row in df_resultados.iterrows():
        # coordenadas (str "lat, lon") -> tuplas float
        lat_pt, lon_pt = map(float, row['Localização do Ponto'].split(', '))
        lat_cx, lon_cx = map(float, row['Localização da Caixa'].split(', '))

        ponto_latlon = (lat_pt, lon_pt)
        caixa_latlon = (lat_cx, lon_cx)

        rota_osrm_coords = row.get('rota_osrm_coords') or []
        rota_postes_coords = row.get('rota_postes_coords') or []

        # respeita a escolha do usuário para desenhar rota de postes
        if not incluir_rota_postes:
            rota_postes_coords = []

        # BOM + pontos dos postes
        bom_dict = None
        postes_pts_coords = row.get('postes_pts_coords') or []
        if rota_postes_coords:
            bom_dict = classificar_postes_por_angulo_e_regra(
                rota_postes_coords, ang_inf=165.0, ang_sup=195.0, max_passantes=4
            )

        nome_base = f"PONTO-{(row.get('Nome do Ponto de Referência') or 'Ponto').replace(' ','_')}__CAIXA-{row.get('Identificador')}"
        png_path = os.path.join(pasta_maps, f"map_{nome_base}.png")
        docx_path = os.path.join(pasta_docs, f"Descritivo_{nome_base}.docx")

        # gera PNG (com presets de qualidade)
        _plot_print_satelital_esri(
            ponto_latlon, caixa_latlon,
            rota_osrm_coords=rota_osrm_coords,
            rota_postes_coords=rota_postes_coords,
            incluir_postes=incluir_postes_pts,
            postes_coords=postes_pts_coords if incluir_postes_pts else None,
            out_png_path=png_path,
            show_osrm=show_osrm,
            show_postes=show_postes,
            dpi=dpi, width_px=width_px, height_px=height_px,
            zoom_override=zoom_override, overlay_labels=overlay_labels
        )

        # gera DOCX a partir do modelo (agora preenchendo QTD_SUPA/BAP/ALCA)
        _criar_docx_descritivo(row, png_path, docx_path, bom_dict=bom_dict, template_path="Modelo de Descritivo.docx")
        docx_paths.append(docx_path)

    # monta ZIP
    zip_path = os.path.join(base_out, "descritivos.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in docx_paths:
            zf.write(p, arcname=os.path.join("docs", os.path.basename(p)))
        for fname in os.listdir(pasta_maps):
            zf.write(os.path.join(pasta_maps, fname), arcname=os.path.join("maps", fname))

    return zip_path
